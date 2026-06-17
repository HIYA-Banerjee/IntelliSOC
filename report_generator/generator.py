import os
import sys
import logging
from datetime import datetime
from typing import List, Dict, Any
from dotenv import load_dotenv

# Import ReportLab for PDF generation
REPORTLAB_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    logger = logging.getLogger("ReportGenerator")
    logger.warning("ReportLab package not found. PDF report generation will use text/markdown fallback.")

# Import python-docx for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ReportGenerator")

class IncidentReportGenerator:
    def __init__(self):
        # On Vercel (or other read-only serverless platforms), write reports to /tmp
        if os.getenv("VERCEL") == "1" or os.environ.get("VERCEL_ENV"):
            self.output_dir = "/tmp/reports"
        else:
            self.output_dir = "reports"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_pdf_report(self, threats: List[Dict[str, Any]], filename: str = None) -> str:
        """
        Generates a highly styled, professional PDF incident report.
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"incident_report_{timestamp}.pdf"
            
        file_path = os.path.join(self.output_dir, filename)
        
        if not REPORTLAB_AVAILABLE:
            logger.warning("ReportLab is not installed. Generating a text/markdown report fallback instead.")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("======================================================================\n")
                f.write("INTELLISOC SECURITY Operations Center - SECURE INCIDENT ADVISORY\n")
                f.write("======================================================================\n\n")
                f.write(f"Generated on:      {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Total Incidents:   {len(threats)}\n\n")
                f.write("1. EXECUTIVE SUMMARY\n")
                f.write("---------------------\n")
                f.write("This report aggregates active network compromises and flow violations analyzed by our ML engines.\n\n")
                f.write("2. INCIDENTS TIMELINE LOG\n")
                f.write("--------------------------\n")
                for t in threats[:15]:
                    f.write(f"- [{t.get('timestamp', '')[:19]}] {t.get('threat_type', '')} ({t.get('severity_level', '')}) | Source: {t.get('source_ip', '')} -> Confidence: {t.get('confidence_score', 0):.0f}%\n")
                f.write("\n3. REMEDIATION ACTIONS & PLAYBOOK\n")
                f.write("----------------------------------\n")
                unique_ts = list(set(t.get("threat_type") for t in threats if t.get("threat_type") != "Normal"))
                for ut in unique_ts:
                    match_t = next((t for t in threats if t.get("threat_type") == ut), {})
                    f.write(f"[*] {ut} Mitigation:\n    {match_t.get('remediation_steps', '')}\n\n")
            return file_path

        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Define premium color palette (Classic navy, slate blue, accent dark red)
        primary_color = colors.HexColor("#0f172a")    # Slate 900
        secondary_color = colors.HexColor("#1e293b")  # Slate 800
        accent_color = colors.HexColor("#ef4444")     # Red 500
        text_color = colors.HexColor("#334155")       # Slate 700
        light_bg = colors.HexColor("#f8fafc")         # Slate 50
        
        # Modify existing styles to avoid collisions
        styles['Normal'].textColor = text_color
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 14
        
        # Create custom paragraph styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Normal'],
            textColor=primary_color,
            fontSize=26,
            leading=32,
            fontName='Helvetica-Bold',
            spaceAfter=6
        )
        
        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            textColor=colors.HexColor("#64748b"),
            fontSize=12,
            leading=16,
            fontName='Helvetica',
            spaceAfter=25
        )
        
        h1_style = ParagraphStyle(
            'ReportH1',
            parent=styles['Normal'],
            textColor=primary_color,
            fontSize=16,
            leading=20,
            fontName='Helvetica-Bold',
            spaceBefore=18,
            spaceAfter=10,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'ReportH2',
            parent=styles['Normal'],
            textColor=secondary_color,
            fontSize=12,
            leading=16,
            fontName='Helvetica-Bold',
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True
        )
        
        callout_style = ParagraphStyle(
            'CalloutText',
            parent=styles['Normal'],
            textColor=colors.HexColor("#1e293b"),
            fontSize=10,
            leading=14,
            fontName='Helvetica-Oblique'
        )
        
        story = []
        
        # --- HEADER / TITLE ---
        story.append(Paragraph("IntelliSOC Incident Security Report", title_style))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Threat Intel Core", subtitle_style))
        story.append(Spacer(1, 15))
        
        # --- EXECUTIVE SUMMARY ---
        story.append(Paragraph("1. Executive Summary", h1_style))
        
        total_threats = len(threats)
        critical_threats = sum(1 for t in threats if t.get("severity_level") == "Critical")
        high_threats = sum(1 for t in threats if t.get("severity_level") == "High")
        medium_threats = sum(1 for t in threats if t.get("severity_level") == "Medium")
        
        summary_text = (
            f"This security intelligence report provides an executive summary and detailed telemetry of network events "
            f"analyzed by the IntelliSOC Threat Detection Platform. During the monitoring window, a total of <b>{total_threats}</b> "
            f"threat incidents were detected. Of these, <b>{critical_threats}</b> were flagged as Critical Severity, and "
            f"<b>{high_threats}</b> were flagged as High Severity. Immediate administrative action and network filtering updates "
            f"are recommended for malicious nodes listed in this report."
        )
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Highlight Callout
        callout_data = [[
            Paragraph(
                f"<b>Threat Advisory:</b> Critical and high severity attacks represent active network compromises "
                f"or high-rate denial of service activities. Attacker IPs have been mapped to threat intelligence indices.",
                callout_style
            )
        ]]
        callout_table = Table(callout_data, colWidths=[500])
        callout_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#fef2f2")),
            ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor("#991b1b")),
            ('BOX', (0,0), (-1,-1), 1.5, colors.HexColor("#fee2e2")),
            ('PADDING', (0,0), (-1,-1), 10),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(callout_table)
        story.append(Spacer(1, 15))
        
        # --- INCIDENT METRICS SUMMARY ---
        story.append(Paragraph("2. Severity Breakdown Statistics", h1_style))
        stats_data = [
            ["Severity Level", "Threat Counts", "Impact Assessment"],
            ["Critical", str(critical_threats), "Immediate denial of service, data exfil, or active botnet communication."],
            ["High", str(high_threats), "Brute force attempts, target port scans, or system probes."],
            ["Medium", str(medium_threats), "Phishing redirections or low-frequency suspicious domain communication."],
            ["Low", str(total_threats - (critical_threats + high_threats + medium_threats)), "Unsupervised anomaly detections and baseline deviations."]
        ]
        
        stats_table = Table(stats_data, colWidths=[110, 90, 300])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), primary_color),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, light_bg]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ('PADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 15))
        
        # --- THREAT TIMELINE ---
        story.append(Paragraph("3. Detailed Threat Incident Log", h1_style))
        
        headers = ["Timestamp", "Source IP", "Destination IP", "Attack Category", "Risk Level", "Conf."]
        table_rows = [headers]
        
        # Limit report list to top 15 rows to prevent oversized PDF documents
        display_threats = threats[:15]
        for t in display_threats:
            # Parse timestamp if it is long
            time_str = t.get("timestamp", "")
            if len(time_str) > 19:
                time_str = time_str[:19]
            
            table_rows.append([
                time_str,
                t.get("source_ip", ""),
                t.get("destination_ip", ""),
                t.get("threat_type", ""),
                t.get("severity_level", ""),
                f"{t.get('confidence_score', 0):.0f}%"
            ])
            
        col_widths = [110, 90, 90, 100, 70, 40]
        timeline_table = Table(table_rows, colWidths=col_widths)
        timeline_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), secondary_color),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, light_bg]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('PADDING', (0,0), (-1,-1), 5),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(timeline_table)
        story.append(Spacer(1, 15))
        
        # --- MITIGATION ACTIONS ---
        story.append(Paragraph("4. Recommended Administrative Actions", h1_style))
        
        # Extract unique threat categories and map actions
        unique_threats = list(set(t.get("threat_type") for t in threats if t.get("threat_type") != "Normal"))
        if not unique_threats:
            story.append(Paragraph("No active remediation required. Continuous packet monitoring recommended.", styles['Normal']))
        else:
            for idx, threat in enumerate(unique_threats):
                story.append(Paragraph(f"<b>4.{idx+1} Remediation for {threat} Incidents</b>", h2_style))
                
                # Fetch matching remediation text
                sample_t = next((t for t in threats if t.get("threat_type") == threat), {})
                remed_text = sample_t.get("remediation_steps", "Apply firewall block rules on targeted interfaces.")
                
                story.append(Paragraph(remed_text, styles['Normal']))
                story.append(Spacer(1, 5))
                
        # Build Document
        doc.build(story)
        logger.info(f"PDF Incident Report compiled at {file_path}")
        return file_path

    def generate_docx_report(self, threats: List[Dict[str, Any]], filename: str = None) -> str:
        """
        Generates an editable MS Word (DOCX) incident report.
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx is not installed. Skipping DOCX generation and returning PDF path instead.")
            # Standard PDF fallback if word generator is unavailable
            return self.generate_pdf_report(threats, filename)
            
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"incident_report_{timestamp}.docx"
            
        file_path = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Cover / Header
        title = doc.add_paragraph()
        run = title.add_run("INTELLISOC SECURITY Operations Center")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run(f"AI-Powered Incident Response Report | Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
        
        doc.add_heading("1. Executive Summary", level=1)
        
        total_threats = len(threats)
        critical_threats = sum(1 for t in threats if t.get("severity_level") == "Critical")
        high_threats = sum(1 for t in threats if t.get("severity_level") == "High")
        medium_threats = sum(1 for t in threats if t.get("severity_level") == "Medium")
        
        summary_text = (
            f"This security intelligence report provides an executive summary and detailed telemetry of network events "
            f"analyzed by the IntelliSOC Threat Detection Platform. During the monitoring window, a total of {total_threats} "
            f"threat incidents were detected. Of these, {critical_threats} were flagged as Critical Severity, and "
            f"{high_threats} were flagged as High Severity. Immediate administrative action and network filtering updates "
            f"are recommended for malicious nodes listed in this report."
        )
        doc.add_paragraph(summary_text)
        
        doc.add_heading("2. Attack Counts and Threat Distribution", level=1)
        
        # Add table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity Level'
        hdr_cells[1].text = 'Alert Counts'
        hdr_cells[2].text = 'Impact Assessment'
        
        stats = [
            ("Critical", str(critical_threats), "Immediate denial of service, active botnet beaconing, or file exfiltrations."),
            ("High", str(high_threats), "Credential brute forcing, ports scans, or network scanning probes."),
            ("Medium", str(medium_threats), "Phishing redirection, suspicious server beaconing."),
            ("Low", str(total_threats - (critical_threats + high_threats + medium_threats)), "Unsupervised anomaly detections and baseline deviations.")
        ]
        
        for level, count, desc in stats:
            row_cells = table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = count
            row_cells[2].text = desc
            
        doc.add_paragraph() # spacing
        
        doc.add_heading("3. Security Alerts Incident Log", level=1)
        
        log_table = doc.add_table(rows=1, cols=6)
        log_table.style = 'Light Shading Accent 1'
        log_hdr = log_table.rows[0].cells
        log_hdr[0].text = 'Timestamp'
        log_hdr[1].text = 'Source IP'
        log_hdr[2].text = 'Destination IP'
        log_hdr[3].text = 'Category'
        log_hdr[4].text = 'Risk Level'
        log_hdr[5].text = 'Conf.'
        
        for t in threats[:15]:
            row_cells = log_table.add_row().cells
            time_str = t.get("timestamp", "")
            row_cells[0].text = time_str[:19] if len(time_str) > 19 else time_str
            row_cells[1].text = t.get("source_ip", "")
            row_cells[2].text = t.get("destination_ip", "")
            row_cells[3].text = t.get("threat_type", "")
            row_cells[4].text = t.get("severity_level", "")
            row_cells[5].text = f"{t.get('confidence_score', 0):.0f}%"
            
        doc.add_paragraph() # spacing
        
        doc.add_heading("4. Mitigation Playbook & Actions", level=1)
        unique_threats = list(set(t.get("threat_type") for t in threats if t.get("threat_type") != "Normal"))
        
        if not unique_threats:
            doc.add_paragraph("No threats were detected in the log. System running within normal parameters.")
        else:
            for threat in unique_threats:
                doc.add_heading(f"Mitigation Plan for {threat} Incidents", level=2)
                sample_t = next((t for t in threats if t.get("threat_type") == threat), {})
                remed_text = sample_t.get("remediation_steps", "Enforce perimeter blocking rules.")
                doc.add_paragraph(remed_text)
                
        doc.save(file_path)
        logger.info(f"DOCX Incident Report compiled at {file_path}")
        return file_path

if __name__ == "__main__":
    generator = IncidentReportGenerator()
    test_threats = [
        {
            "timestamp": "2026-06-14 13:00:00",
            "source_ip": "198.51.100.42",
            "destination_ip": "192.168.1.50",
            "threat_type": "DDoS",
            "severity_level": "Critical",
            "confidence_score": 98.5,
            "risk_score": 94.0,
            "remediation_steps": "Apply firewall blocks for IP 198.51.100.42"
        },
        {
            "timestamp": "2026-06-14 13:05:00",
            "source_ip": "203.0.113.195",
            "destination_ip": "192.168.1.20",
            "threat_type": "Brute Force",
            "severity_level": "High",
            "confidence_score": 88.0,
            "risk_score": 75.0,
            "remediation_steps": "Rate-limit port 22 connections."
        }
    ]
    generator.generate_pdf_report(test_threats, "test_report.pdf")
    generator.generate_docx_report(test_threats, "test_report.docx")
